# 💻 My_Python_Projects

#1. 🕒   Digital watch with alarm :

import tkinter as tk
from time import strftime
import threading
import time
from termcolor import colored  # Importăm termcolor pentru a schimba culoarea textului

# Fereastra pentru ceas
root = tk.Tk()
root.title('Digital Clock with Alarm')

# Eticheta pentru ceas
clock_label = tk.Label(root, font=('Helvetica', 48), bg='black', fg='cyan')
clock_label.pack(anchor='center', fill='both', expand=True)

# Variabila care va ține ora alarmei
alarm_time = None

# Funcția care actualizează timpul pe ceas
def update_time():
    current_time = strftime('%H:%M:%S')  # Obține ora curentă
    clock_label.config(text=current_time)  # Actualizează ceasul

    # Verifică dacă ora curentă corespunde cu ora alarmei
    if alarm_time and current_time[:5] == alarm_time:
        print(colored(f"A sunat alarma la {alarm_time}!", 'red'))  # Afișează mesajul roșu în consolă
        
    clock_label.after(1000, update_time)  # Actualizează ceasul la fiecare secundă

# Funcția pentru setarea alarmei
def set_alarm():
    global alarm_time
    alarm_input = input("Setează ora alarmei (format HH:MM): ")

    try:
        # Verifică dacă inputul este într-un format corect
        time.strptime(alarm_input, '%H:%M')  # Verifică dacă formatul e corect
        alarm_time = alarm_input  # Salvează ora alarmei
        print(f"Alarma este setată la: {alarm_time}")
    except ValueError:
        print("Format invalid. Te rog să introduci ora în formatul HH:MM.")

# Setează alarma într-un thread separat
def alarm_thread_function():
    while True:
        set_alarm()
        time.sleep(1)  # Evită blocarea programului

# Pornește funcția de actualizare a timpului și setarea alarmei în paralel
thread = threading.Thread(target=alarm_thread_function)
thread.daemon = True
thread.start()

update_time()

root.mainloop()


print("&" * 20)


#2.  🗺️  Create map / search location using  Python:

import folium
from geopy.geocoders import Nominatim
import webbrowser
import os

location_name = input('Enter a location:   ')

geolocator = Nominatim(user_agent='geoapi')
location = geolocator.geocode(location_name)

if location:
    # Create a map centered on user's location
    latitude = location.latitude
    longitude = location.longitude
    clcoding = folium.Map(location=[latitude, longitude], zoom_start=12)

    marker = folium.Marker([latitude, longitude], popup=location_name)
    marker.add_to(clcoding)

    # Save the map to an HTML file and open it in the default web browser
    map_file = 'location_map.html'
    clcoding.save(map_file)
    webbrowser.open('file://' + os.path.realpath(map_file))
else:
    print('Location not found. Please try again.')

print("&" * 20)


#3.  🔺  Python code for PASCAL's TRIANGLE:

def printPascal(N):
    arr = [1]  # Initialize the first row
    print("Pascal's triangle with", N, "rows:")

    for i in range(N):
        print("Row", i + 1, end=" : ")

        # Print the current row
        for val in arr:
            print(val, end=" ")
        print()  # Newline after each row

        # Generate the next row
        temp = [1]  # Start the next row with 1
        for j in range(len(arr) - 1):
            temp.append(arr[j] + arr[j + 1])  # Sum of adjacent elements
        temp.append(1)  # End the row with 1
        arr = temp  # Update the current row for the next iteration


# Get user input and print Pascal's triangle
N = int(input('Enter the number of rows for Pascal\'s triangle: '))
printPascal(N)

print("&" * 20)


#4.  🏳️‍🌈  Myanmar flag using Python:

import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
import numpy as np

fig, ax = plt.subplots(figsize=(8, 5))
ax.fill_between([0, 3], 2, 3, color="#FED100")
ax.fill_between([0, 3], 1, 2, color="#34B233")
ax.fill_between([0, 3], 0, 1, color="#EA2839")

def draw_star(center_x, center_y, radius, color, rotation_deg):
    points = []
    for i in range(10):
        angle = (i * 36 + rotation_deg) * (np.pi / 180)
        r = radius if i % 2 == 0 else radius / 2
        x = center_x + r * np.cos(angle)
        y = center_y + r * np.sin(angle)
        points.append((x, y))
    polygon = Polygon(points, closed=True, color=color)
    ax.add_patch(polygon)

draw_star(1.5, 1.5, 0.6, "white", rotation_deg=-55)
ax.set_xlim(0, 3)
ax.set_ylim(0, 3)
ax.axis("off")
plt.show()
print("Happy independence Day Myanmar! ")

print("&" * 20)


#5.   🟦🟨🟥  Romanian flag using Python:

import matplotlib.pyplot as plt

# Create a figure and axis
fig, ax = plt.subplots(figsize=(9, 5))

# Define the width of each stripe
stripe_width = 1 / 3

# Draw the blue stripe (left)
ax.fill_betweenx([0, 1], 0, stripe_width, color='#002B7F')  # Blue color

# Draw the yellow stripe (center)
ax.fill_betweenx([0, 1], stripe_width, 2 * stripe_width, color='#FCD116')  # Yellow color

# Draw the red stripe (right)
ax.fill_betweenx([0, 1], 2 * stripe_width, 1, color='#CE1126')  # Red color

# Set limits and remove axes
ax.set_xlim(0, 1)
ax.set_ylim(0, 1)
ax.axis('off')

# Display the flag
plt.title("La multi ani Romania!", fontsize=15)
plt.show()

print("&" * 20)



#6.  📄➡️📝  CONVERSION PDF TO WORD


import os
import PyPDF2
from docx import Document


# Function to search for the PDF file using os.walk()
def search_pdf_file(filename, start_dir="/"):
    for root, dirs, files in os.walk(start_dir):
        if filename in files:
            return os.path.join(root, filename)
    return None


# Specify the filename you're searching for
filename = "118567_6513eb3d888fa9.83734900.pdf"

# Search for the PDF file in the specified directory (e.g., C:/ for Windows or / for Unix-like systems)
pdf_path = search_pdf_file(filename, start_dir="C:/")  # Or "/" for Linux/Mac

if pdf_path:
    print(f"File found at: {pdf_path}")

    # Create a new Word document
    doc = Document()

    # Open the found PDF file
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        # Loop through each page of the PDF
        for page in pdf_reader.pages:
            text = page.extract_text()  # Extract text from each page

            # If text is found, add it to the Word document as a paragraph
            if text:
                doc.add_paragraph(text)
            else:
                doc.add_paragraph("No text found on this page.")

    # Save the Word document
    word_filename = filename.replace('.pdf', '.docx')
    doc.save(word_filename)
    print(f"Conversion complete. The document is saved as '{word_filename}'.")

else:
    print("PDF file not found.")

print("&" * 20)



#7.  🌍📡🔍  IP Adress Information Using Python:

#Without VPN  

import urllib.request as urllib2
import json

def get_public_ip():
    """Obține IP-ul public al utilizatorului folosind un serviciu online."""
    try:
        response = urllib2.urlopen("https://api64.ipify.org?format=json")
        data = json.loads(response.read().decode("utf-8"))
        return data["ip"]
    except Exception as e:
        print(f"Could not retrieve public IP: {e}")
        return None

while True:
    ip = input("What is your target IP (or type 'my' to get your own IP): ").strip()

    # Dacă utilizatorul introduce 'my', află IP-ul public automat
    if ip.lower() == "my":
        ip = get_public_ip()
        if ip:
            print(f"Your public IP is: {ip}")
        else:
            print("Could not determine your public IP.")
            continue

    # Verificăm dacă utilizatorul a introdus un IP valid
    if not ip:
        print("Please enter a valid IP address.")
        continue

    url = f"http://ip-api.com/json/{ip}"

    try:
        response = urllib2.urlopen(url)
        data = response.read()
        values = json.loads(data)

        # Verificăm dacă API-ul a returnat succes
        if values.get("status") == "success":
            print("\n🌍 IP Information:")
            print(f"📍 IP: {values.get('query', 'N/A')}")
            print(f"🏙️ City: {values.get('city', 'N/A')}")
            print(f"🗺️ Country: {values.get('country', 'N/A')}")
            print(f"📡 ISP: {values.get('isp', 'N/A')}")
            print(f"🌎 Region: {values.get('regionName', 'N/A')}")
            print(f"⏰ Timezone: {values.get('timezone', 'N/A')}\n")
        else:
            print(f"❌ Error: {values.get('message', 'Unknown error')}")

    except Exception as e:
        print(f"⚠️ An error occurred: {e}")

    break  # Oprește execuția după o singură rulare

print("&" * 20)

#8.  🔑🔐  Alegere parola random (cel putin :2 litere mici, 2 litere mari, 2 simboluri, 2 numere) si cu verificare complexitate:

import random
import string

# Definirea caracterelor
lower = string.ascii_lowercase
upper = string.ascii_uppercase
numbers = string.digits
symbols = "!@#$%^&*()_-+=?><[]"

all_chars = lower + upper + numbers + symbols

# Funcție pentru a verifica complexitatea parolei
def check_password_complexity(password):
    has_lower = sum(1 for char in password if char in lower) >= 2
    has_upper = sum(1 for char in password if char in upper) >= 2
    has_number = sum(1 for char in password if char in numbers) >= 2
    has_symbol = sum(1 for char in password if char in symbols) >= 2

    if has_lower and has_upper and has_number and has_symbol:
        return True
    else:
        return False

# Generarea parolei
length = int(input("Please enter password length: "))

# Asigură-te că parola generată are suficient de multe caractere pentru cerințele impuse
while length < 8:  # minim 8 caractere pentru a satisface cerințele
    print("Parola trebuie să aibă cel puțin 8 caractere.")
    length = int(input("Please enter password length: "))

password = "".join(random.sample(all_chars, length))

# Verificarea complexității
if check_password_complexity(password):
    print("Generated password: ", password)
    print("The password meets the complexity requirements!")
else:
    print("Generated password: ", password)
    print("The password does NOT meet the complexity requirements.")



print("*" * 20)


#9. 🛍️🏷️🛒  Meniu magazin cu introducere articole in inventar, modificare inventar si cumparare.
  
# --------------------------------------
#           Meniu Magazin
# --------------------------------------
# 1. Introdu inventarul
# 2. Șterge sau modifică produs din inventar
# 3. Vizualizează produse
# 4. Adaugă produs în coș
# 5. Vizualizează coșul
# 6. Finalizează comanda
# 7. Iesiti
# --------------------------------------
# Alege o opțiune: 


from colorama import init, Fore, Style

# Inițializare colorama
init(autoreset=True)

# Coșul de cumpărături (inițial gol)
cart = {}

# Lista de produse adaugate in inventar si disponibile cu stoc
products = {}

# Funcția pentru a introduce inventarul
def initialize_inventory():
    global products
    print(Fore.CYAN + "Introdu produsele disponibile în magazin.")
    
    while True:
        name = input(Fore.YELLOW + "Nume produs (sau 'stop' pentru a termina): ").strip()
        if name.lower() == "stop":
            break
        
        try:
            price = float(input(Fore.YELLOW + f"Pret pentru {name}: "))
            stock = int(input(Fore.YELLOW + f"Stoc pentru {name}: "))
            products[name] = {"price": price, "stock": stock}
            print(Fore.GREEN + f"Produsul {name} a fost adăugat cu succes!\n")
        except ValueError:
            print(Fore.RED + "Date invalide! Introdu un pret și un stoc valide.")

# Funcția pentru a șterge sau modifica produse
def modify_product():
    if not products:
        print(Fore.RED + "Nu sunt produse în inventar. Te rog să introduci inventarul mai întâi.")
        return

    print(Fore.GREEN + "\nProduse disponibile:")
    for index, (product, info) in enumerate(products.items(), start=1):
        print(Fore.YELLOW + f"{index}. {product} - {info['price']} RON (Stoc: {info['stock']})")
    
    product_name = input(Fore.YELLOW + "\nIntrodu numele produsului pe care vrei să-l modifici sau 'stop' pentru a renunța: ").strip()
    if product_name.lower() == "stop":
        return

    if product_name in products:
        print(Fore.GREEN + f"Modifică produsul {product_name}:")
        sub_option = input(Fore.YELLOW + "Vrei să modifici cantitatea sau pretul? (cantitate/pret): ").strip().lower()

        if sub_option == "cantitate":
            try:
                quantity = int(input(Fore.YELLOW + f"Introdu cantitatea totală pe care vrei să o ai pentru {product_name}: "))
                if quantity >= 0:
                    # Actualizez stocul produsului direct la cantitatea introdusă
                    products[product_name]["stock"] = quantity
                    print(Fore.GREEN + f"Stocul pentru {product_name} a fost actualizat la {quantity}.")
                else:
                    print(Fore.RED + "Cantitatea nu poate fi negativă.")
            except ValueError:
                print(Fore.RED + "Te rog să introduci o cantitate validă.")
        elif sub_option == "pret":
            try:
                new_price = float(input(Fore.YELLOW + f"Introdu noul pret pentru {product_name}: "))
                if new_price > 0:
                    products[product_name]["price"] = new_price
                    print(Fore.GREEN + f"Prețul pentru {product_name} a fost actualizat la {new_price} RON.")
                else:
                    print(Fore.RED + "Pretul nu poate fi mai mic de 0.")
            except ValueError:
                print(Fore.RED + "Te rog să introduci un pret valid.")
        else:
            print(Fore.RED + "Opțiune invalidă. Te rog să alegi 'cantitate' sau 'pret'.")
    else:
        print(Fore.RED + "Produsul nu există în inventar.")

# Funcția pentru a vizualiza produsele disponibile
def show_products():
    if not products:
        print(Fore.RED + "Nu sunt produse în inventar. Te rog să introduci inventarul mai întâi.")
        return
    
    print(Fore.GREEN + "\nProduse disponibile:")
    for index, (product, info) in enumerate(products.items(), start=1):
        print(Fore.YELLOW + f"{index}. {product} - {info['price']} RON (Stoc: {info['stock']})")

# Funcția pentru a adăuga produse în coș
def add_to_cart():
    if not products:
        print(Fore.RED + "Nu sunt produse în inventar. Te rog să introduci inventarul mai întâi.")
        return
    
    show_products()
    
    try:
        choice = int(input(Fore.YELLOW + "\nAlege numărul produsului pentru a-l adăuga în coș: "))
        if 1 <= choice <= len(products):
            product_name = list(products.keys())[choice - 1]
            quantity = int(input(Fore.YELLOW + f"Câte {product_name} vrei să adaugi în coș? "))

            # Verificăm dacă există suficiente produse în stoc
            if quantity <= products[product_name]["stock"]:
                if product_name in cart:
                    cart[product_name] += quantity
                else:
                    cart[product_name] = quantity

                # Scădem cantitatea din stoc
                products[product_name]["stock"] -= quantity
                print(Fore.GREEN + f"Produsul {product_name} a fost adăugat în coș!")
            else:
                print(Fore.RED + f"Nu sunt suficiente produse în stoc pentru {product_name}. Mai sunt doar {products[product_name]['stock']} disponibile.")
        else:
            print(Fore.RED + "Opțiune invalidă.")
    except ValueError:
        print(Fore.RED + "Te rog să introduci un număr valid.")

# Funcția pentru a vizualiza coșul de cumpărături
def view_cart():
    if cart:
        print(Fore.GREEN + "\nCoșul tău de cumpărături:")
        total_price = 0
        for product, quantity in cart.items():
            price = products[product]["price"]
            total_price += price * quantity
            print(Fore.YELLOW + f"{product} - {quantity} x {price} RON")
        print(Fore.YELLOW + f"\nTotal: {total_price} RON")
    else:
        print(Fore.RED + "Coșul tău este gol.")

# Funcția pentru a finaliza comanda
def checkout():
    if cart:
        print(Fore.GREEN + "\nMulțumim pentru cumpărături!")
        total_price = 0
        for product, quantity in cart.items():
            price = products[product]["price"]
            total_price += price * quantity
            print(Fore.YELLOW + f"{product} - {quantity} x {price} RON")
        print(Fore.YELLOW + f"\nTotal de plată: {total_price} RON")

        # Golește coșul după ce comanda este finalizată
        cart.clear()
        print(Fore.GREEN + "\nComanda a fost finalizată. La revedere!")
    else:
        print(Fore.RED + "Coșul este gol. Nu poți finaliza comanda fără produse.")

# Funcția care arată meniul principal
def show_menu():
    print(Fore.YELLOW + "\n--------------------------------------")
    print(Fore.YELLOW + "          Meniu Magazin")
    print(Fore.YELLOW + "--------------------------------------")
    print(Fore.GREEN + "1. Introdu inventarul")
    print(Fore.GREEN + "2. Șterge sau modifică produs din inventar")
    print(Fore.GREEN + "3. Vizualizează produse")
    print(Fore.GREEN + "4. Adaugă produs în coș")
    print(Fore.GREEN + "5. Vizualizează coșul")
    print(Fore.GREEN + "6. Finalizează comanda")
    print(Fore.GREEN + "7. Ieși")
    print(Fore.YELLOW + "--------------------------------------")

# Funcția principală
def main():
    while True:
        show_menu()

        try:
            option = int(input(Fore.YELLOW + "Alege o opțiune: "))

            if option == 1:
                initialize_inventory()
            elif option == 2:
                modify_product()
            elif option == 3:
                show_products()
            elif option == 4:
                add_to_cart()
            elif option == 5:
                view_cart()
            elif option == 6:
                checkout()
                break  # După ce finalizăm comanda, ieșim din aplicație
            elif option == 7:
                print(Fore.GREEN + "Ieșire din aplicație. La revedere!")
                break
            else:
                print(Fore.RED + "Opțiune invalidă. Te rog să alegi 1, 2, 3, 4, 5, 6 sau 7.")
        except ValueError:
            print(Fore.RED + "Te rog să introduci un număr valid.\n")


if __name__ == "__main__":
    main()

print("*" * 20)
