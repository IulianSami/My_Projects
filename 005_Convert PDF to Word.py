#5.  📄➡️📝  CONVERT PDF TO WORD


import os
import PyPDF2
from docx import Document


# Function to search for the PDF file using os.walk()
def search_pdf_file(filename, start_dir="/"):
    for root, dirs, files in os.walk(start_dir):
        if filename in files:
            return os.path.join(root, filename)
    return None


# Specify the filename you're searching for
filename = "118567_6513eb3d888fa9.83734900.pdf"

# Search for the PDF file in the specified directory (e.g., C:/ for Windows or / for Unix-like systems)
pdf_path = search_pdf_file(filename, start_dir="C:/")  # Or "/" for Linux/Mac

if pdf_path:
    print(f"File found at: {pdf_path}")

    # Create a new Word document
    doc = Document()

    # Open the found PDF file
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        # Loop through each page of the PDF
        for page in pdf_reader.pages:
            text = page.extract_text()  # Extract text from each page

            # If text is found, add it to the Word document as a paragraph
            if text:
                doc.add_paragraph(text)
            else:
                doc.add_paragraph("No text found on this page.")

    # Save the Word document
    word_filename = filename.replace('.pdf', '.docx')
    doc.save(word_filename)
    print(f"Conversion complete. The document is saved as '{word_filename}'.")

else:
    print("PDF file not found.")